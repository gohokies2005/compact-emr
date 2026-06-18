"""Package 2 (keystone plan) tests: native text-readers in the OCR start handler.

Covers: .txt decode + the /pages post shape; .docx extraction (fixture built with the LOCAL
python-docx — see conftest.py); the legacy-.doc best-effort ladder including the OLE flag
note; extension-not-contentType keying; and the guards around the native branch (hasPages
skip, intake/ prefix untouched, oversize flag, pdf still rides Textract).

No network, no AWS: _resolve_document / _post_pages_to_api / _post_failed_read_attempt /
s3 / textract are all monkeypatched.
"""
import io
import json

import pytest

import handler  # workers/ocr/handler.py — sys.path arranged by conftest.py


def _event(key: str, bucket: str = "phi-bucket") -> dict:
    """EventBridge ObjectCreated shape — note there is NO contentType anywhere in it: the
    native branch can only key on the s3-key extension (the intakes.ts effectiveContentType
    lesson; the declared upload contentType is application/octet-stream anyway)."""
    return {"detail": {"bucket": {"name": bucket}, "object": {"key": key}}}


class _S3Stub:
    def __init__(self, data: bytes):
        self.data = data
        self.calls: list[tuple[str, str]] = []

    def get_object(self, Bucket: str, Key: str) -> dict:  # noqa: N803 — boto3 casing
        self.calls.append((Bucket, Key))
        return {"Body": io.BytesIO(self.data)}


class _TextractBomb:
    """Any Textract call on a native-read path is a test failure."""

    def __getattr__(self, name: str):
        raise AssertionError(f"textract.{name} must not be called on the native-read path")


class _TextractStub:
    def __init__(self):
        self.calls: list[dict] = []

    def start_document_text_detection(self, **kwargs) -> dict:
        self.calls.append(kwargs)
        return {"JobId": "JOB-1"}


@pytest.fixture
def rig(monkeypatch):
    """start_handler wired for a resolvable cases/ document with no pages; records all posts."""
    state = {
        "pages": [],
        "failed": [],
        "doc": {"documentId": "DOC-1", "hasPages": False},
    }
    monkeypatch.setattr(handler, "_resolve_document", lambda key: state["doc"])
    monkeypatch.setattr(
        handler, "_post_pages_to_api",
        lambda doc_id, pages, count: state["pages"].append((doc_id, pages, count)),
    )
    monkeypatch.setattr(
        handler, "_post_failed_read_attempt",
        lambda doc_id, status, job_id, error_message=None: state["failed"].append((doc_id, status, job_id, error_message)),
    )
    monkeypatch.setattr(handler, "textract", _TextractBomb())

    def with_bytes(data: bytes) -> _S3Stub:
        stub = _S3Stub(data)
        monkeypatch.setattr(handler, "s3", stub)
        return stub

    state["with_bytes"] = with_bytes
    return state


def _docx_bytes() -> bytes:
    """Build the .docx fixture with the LOCAL python-docx (paragraph, table, paragraph)."""
    import docx as docx_pkg

    d = docx_pkg.Document()
    d.add_paragraph("Chief complaint: chronic lumbar pain.")
    table = d.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Medication"
    table.rows[0].cells[1].text = "Dose"
    table.rows[1].cells[0].text = "Meloxicam"
    table.rows[1].cells[1].text = "15 mg"
    d.add_paragraph("Assessment: degenerative disc disease.")
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


# ===== .txt =====

def test_txt_posts_single_page_native_text(rig):
    content = "Sleep study: AHI 36.4 events/hour. Severe OSA.\nCPAP initiated."
    s3 = rig["with_bytes"](content.encode("utf-8"))
    result = handler.start_handler(_event("cases/C1/u1-records.txt"), None)

    assert result["native"] == "txt"
    assert result["method"] == "native_text"
    assert result["started"] == []
    assert rig["failed"] == []
    [(doc_id, pages, count)] = rig["pages"]
    assert doc_id == "DOC-1"
    assert count == 1
    assert pages == [{"pageNumber": 1, "text": content, "confidence": None}]
    assert s3.calls == [("phi-bucket", "cases/C1/u1-records.txt")]


def test_txt_extension_keying_is_case_insensitive(rig):
    rig["with_bytes"](b"VA rating decision narrative, service connection granted.")
    result = handler.start_handler(_event("cases/C1/u2-NOTES.TXT"), None)
    assert result["method"] == "native_text"
    assert len(rig["pages"]) == 1


def test_txt_utf8_bom_is_stripped(rig):
    rig["with_bytes"](b"\xef\xbb\xbf" + "Tinnitus reported constant and bilateral.".encode("utf-8"))
    handler.start_handler(_event("cases/C1/u3-note.txt"), None)
    [(_, pages, _)] = rig["pages"]
    assert pages[0]["text"].startswith("Tinnitus")  # no ﻿ residue


def test_txt_utf16_bom_decodes(rig):
    content = "PTSD screening résumé — Notepad saves UTF-16."
    rig["with_bytes"](content.encode("utf-16"))  # BOM-prefixed UTF-16 LE
    handler.start_handler(_event("cases/C1/u4-note.txt"), None)
    [(_, pages, _)] = rig["pages"]
    assert pages[0]["text"] == content


def test_txt_form_feed_splits_pages(rig):
    rig["with_bytes"]("Page one body.\fPage two body.".encode("utf-8"))
    result = handler.start_handler(_event("cases/C1/u5-multi.txt"), None)
    [(_, pages, count)] = rig["pages"]
    assert result["pages"] == 2
    assert count == 2
    assert [p["pageNumber"] for p in pages] == [1, 2]
    assert pages[0]["text"] == "Page one body."
    assert pages[1]["text"] == "Page two body."


def test_txt_empty_file_still_posts_a_page_for_the_classifier(rig):
    rig["with_bytes"](b"")
    handler.start_handler(_event("cases/C1/u6-empty.txt"), None)
    [(_, pages, _)] = rig["pages"]
    assert pages == [{"pageNumber": 1, "text": "", "confidence": None}]
    assert rig["failed"] == []  # word-count gating happens server-side in the SAME /pages path


# ===== .html (E4 — VA Rated-Disabilities / Blue Button exports, tags stripped, no Textract) =====

def test_html_strips_tags_and_posts_native_text(rig):
    html = (
        b"<html><head><style>td{color:red}</style></head><body>"
        b"<h2>Rated Disabilities</h2>"
        b"<table><tr><td>PTSD</td><td>70%</td><td>Service Connected</td></tr>"
        b"<tr><td>Tinnitus</td><td>10%</td><td>Service Connected</td></tr></table>"
        b"<script>track();</script></body></html>"
    )
    rig["with_bytes"](html)
    result = handler.start_handler(_event("cases/C1/u7-Rated_Disabilities.html"), None)

    assert result["native"] == "html"
    assert result["method"] == "native_html"
    assert result["started"] == []
    assert rig["failed"] == []
    [(_, pages, _)] = rig["pages"]
    text = pages[0]["text"]
    assert "PTSD" in text and "70%" in text and "Service Connected" in text
    assert "Tinnitus" in text
    assert "track()" not in text and "color:red" not in text  # script/style dropped


def test_html_extension_is_case_insensitive_and_htm_supported(rig):
    rig["with_bytes"](b"<html><body><p>Service connection granted.</p></body></html>")
    result = handler.start_handler(_event("cases/C1/u8-DECISION.HTM"), None)
    assert result["method"] == "native_html"
    [(_, pages, _)] = rig["pages"]
    assert "Service connection granted." in pages[0]["text"]


def test_html_with_no_readable_text_flags_for_rn(rig):
    rig["with_bytes"](b"<html><head><script>var x=1;</script></head><body></body></html>")
    result = handler.start_handler(_event("cases/C1/u9-empty.html"), None)
    assert result.get("flaggedForRn") is True
    assert rig["pages"] == []          # never posts an empty page when the strip yields nothing
    assert len(rig["failed"]) == 1     # actionable RN flag, never silent


# ===== .docx =====

def test_docx_extracts_paragraphs_and_tables_in_document_order(rig):
    rig["with_bytes"](_docx_bytes())
    result = handler.start_handler(_event("cases/C1/u7-chart.docx"), None)

    assert result["method"] == "native_docx"
    [(_, pages, _)] = rig["pages"]
    text = pages[0]["text"]
    i_para1 = text.index("Chief complaint: chronic lumbar pain.")
    i_table = text.index("Medication | Dose")
    i_row2 = text.index("Meloxicam | 15 mg")
    i_para2 = text.index("Assessment: degenerative disc disease.")
    assert i_para1 < i_table < i_row2 < i_para2


def test_docx_corrupt_bytes_flag_with_actionable_note(rig):
    rig["with_bytes"](b"this is not a zip archive at all")
    result = handler.start_handler(_event("cases/C1/u8-broken.docx"), None)

    assert result["flaggedForRn"] is True
    assert rig["pages"] == []
    [(doc_id, status, job_id, note)] = rig["failed"]
    assert (doc_id, status, job_id) == ("DOC-1", "NATIVE_UNREADABLE", "native-read")
    assert ".docx could not be parsed" in note
    assert "summarize manually" in note


# ===== legacy .doc ladder =====

def test_doc_mislabeled_docx_reads_via_python_docx(rig):
    rig["with_bytes"](_docx_bytes())  # a real .docx wearing a .doc name
    result = handler.start_handler(_event("cases/C1/u9-old.doc"), None)
    assert result["method"] == "native_docx"
    [(_, pages, _)] = rig["pages"]
    assert "Chief complaint: chronic lumbar pain." in pages[0]["text"]


def test_doc_rtf_header_strips_to_text(rig):
    rtf = (
        b"{\\rtf1\\ansi\\deff0{\\fonttbl{\\f0 Calibri;}}\\f0\\fs22 "
        b"Veteran reports knee pain since 2009.\\par Gait antalgic on exam.\\par}"
    )
    rig["with_bytes"](rtf)
    result = handler.start_handler(_event("cases/C1/u10-legacy.doc"), None)
    assert result["method"] == "native_text"
    [(_, pages, _)] = rig["pages"]
    assert "Veteran reports knee pain since 2009." in pages[0]["text"]
    assert "Gait antalgic on exam." in pages[0]["text"]
    assert "\\rtf" not in pages[0]["text"]


def test_doc_mostly_printable_bytes_read_as_plain_text(rig):
    body = b"Progress note: veteran ambulates without assistive device.\r\nGait steady.\r\n" * 3
    rig["with_bytes"](body)
    result = handler.start_handler(_event("cases/C1/u11-renamed.doc"), None)
    assert result["method"] == "native_text"
    [(_, pages, _)] = rig["pages"]
    assert "ambulates without assistive device" in pages[0]["text"]


def test_doc_genuine_ole_flags_with_actionable_legacy_note(rig):
    ole = handler._OLE_MAGIC + bytes([0, 1, 2, 3]) * 128
    rig["with_bytes"](ole)
    result = handler.start_handler(_event("cases/C1/u12-word97.doc"), None)

    assert result["flaggedForRn"] is True
    assert rig["pages"] == []
    [(doc_id, status, job_id, note)] = rig["failed"]
    assert status == "LEGACY_DOC"
    assert job_id == "native-read"
    assert "legacy .doc format" in note
    assert "PDF/docx" in note
    assert "summarize manually" in note


def test_doc_binary_garbage_flags_unreadable(rig):
    rig["with_bytes"](bytes(range(256)) * 8)
    result = handler.start_handler(_event("cases/C1/u13-mystery.doc"), None)
    assert result["flaggedForRn"] is True
    [(_, status, _, note)] = rig["failed"]
    assert status == "NATIVE_UNREADABLE"
    assert "unreadable .doc" in note


# ===== Layer 1: native PDF text-layer extraction (pypdf, BEFORE Textract) =====
# A born-digital VA Blue Button dump (a real EMBEDDED text layer on every page) is read DIRECTLY with
# pypdf and POSTed through the SAME /pages pipeline as Textract — Lozano's 2,294-page dump choked
# Textract image-OCR and stored NO pages. A true image-only scan (thin/empty text layer) FALLS THROUGH
# to Textract exactly as before. These tests use a real hand-built PDF for the text-layer + scanned
# cases (exercising the actual pypdf open + extract), and a mocked PdfReader for the >2000-page
# batching + error-fall-through cases.


def _make_text_pdf(page_texts: list[str]) -> bytes:
    """Build a valid multi-page PDF with correct xref byte-offsets and a REAL text layer per page
    (a content stream with a `(...) Tj` show-text op pypdf extracts). No external dependency — the
    born-digital fixture the Layer-1 probe must accept."""
    n = len(page_texts)
    page_obj_nums = [3 + 2 * i for i in range(n)]
    content_obj_nums = [4 + 2 * i for i in range(n)]
    font_obj_num = 3 + 2 * n
    parts: list[bytes] = [b"%PDF-1.4\n"]
    offsets: dict[int, int] = {}

    def add(num: int, body: bytes) -> None:
        offsets[num] = sum(len(p) for p in parts)
        parts.append(("%d 0 obj" % num).encode() + body + b"endobj\n")

    add(1, b"<</Type/Catalog/Pages 2 0 R>>")
    kids = " ".join("%d 0 R" % p for p in page_obj_nums)
    add(2, ("<</Type/Pages/Kids[%s]/Count %d>>" % (kids, n)).encode())
    for i in range(n):
        stream = ("BT /F1 14 Tf 72 700 Td (%s) Tj ET" % page_texts[i]).encode()
        add(page_obj_nums[i], ("<</Type/Page/Parent 2 0 R/MediaBox[0 0 612 792]/Resources<</Font<</F1 %d 0 R>>>>/Contents %d 0 R>>" % (font_obj_num, content_obj_nums[i])).encode())
        add(content_obj_nums[i], ("<</Length %d>>stream\n" % len(stream)).encode() + stream + b"\nendstream")
    add(font_obj_num, b"<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>")
    xref_pos = sum(len(p) for p in parts)
    total = font_obj_num + 1
    xref = ["xref", "0 %d" % total, "0000000000 65535 f "]
    for num in range(1, total):
        xref.append("%010d 00000 n " % offsets[num])
    parts.append(("\n".join(xref) + "\n").encode())
    parts.append(("trailer<</Root 1 0 R/Size %d>>\nstartxref\n%d\n%%%%EOF" % (total, xref_pos)).encode())
    return b"".join(parts)


def _make_blank_pdf(n: int = 1) -> bytes:
    """An n-page PDF with NO text layer (blank pages) — the image-only-scan stand-in. pypdf opens it
    fine but extract_text() yields '' per page, so the probe is thin → Textract fall-through."""
    from pypdf import PdfWriter

    w = PdfWriter()
    for _ in range(n):
        w.add_blank_page(width=200, height=200)
    buf = io.BytesIO()
    w.write(buf)
    return buf.getvalue()


def test_native_pdf_text_layer_posts_per_page_and_never_calls_textract(rig):
    """A digital text-layer PDF is read by pypdf and POSTed per page through the SAME /pages upsert —
    NO Textract (rig's textract is the _TextractBomb)."""
    # Realistic Blue Button page density (hundreds of chars/page) so the probe clears the text-layer
    # floor — a real digital VA export is text-dense, not three short phrases.
    rig["with_bytes"](_make_text_pdf([
        "Page one sleep study report. AHI 36.4 events per hour, severe obstructive sleep apnea confirmed. Nadir oxygen saturation 81 percent. Continuous positive airway pressure therapy recommended at this visit.",
        "Page two treatment plan. CPAP initiated nightly use, pressure 11 cm H2O. Veteran reports daytime somnolence, Epworth Sleepiness Scale score 16. Compliance to be reviewed at the ninety day follow-up appointment.",
        "Page three follow-up note. Sleep study ordered for titration. Service connection sought for obstructive sleep apnea secondary to post traumatic stress disorder. Records forwarded to the rating specialist for review.",
    ]))
    result = handler.start_handler(_event("cases/C1/u30-bluebutton.pdf"), None)

    assert result["native"] == "pdf"
    assert result["method"] == "native_pdf_text"
    assert result["via"] == "pypdf"
    assert result["started"] == []
    assert rig["failed"] == []
    [(doc_id, pages, count)] = rig["pages"]
    assert doc_id == "DOC-1"
    assert count == 3                      # the TRUE pdf page count is sent as documentPageCount
    assert [p["pageNumber"] for p in pages] == [1, 2, 3]
    assert "AHI 36.4" in pages[0]["text"]
    assert "CPAP initiated" in pages[1]["text"]
    assert "follow-up note" in pages[2]["text"]
    assert all(p["confidence"] is None for p in pages)  # native read has no OCR confidence


def test_scanned_pdf_with_no_text_layer_falls_through_to_textract(rig, monkeypatch):
    """A textless (image-only) PDF: pypdf opens it but the probe is thin → Textract starts, exactly
    as before Layer 1. The native branch posts NOTHING."""
    stub = _TextractStub()
    monkeypatch.setattr(handler, "textract", stub)
    monkeypatch.setenv("COMPLETION_SNS_TOPIC_ARN", "arn:aws:sns:us-east-1:0:topic")
    monkeypatch.setenv("TEXTRACT_SNS_ROLE_ARN", "arn:aws:iam::0:role/textract")
    rig["with_bytes"](_make_blank_pdf(3))

    result = handler.start_handler(_event("cases/C1/u31-scan.pdf"), None)

    assert result == {"started": [{"documentId": "DOC-1", "jobId": "JOB-1"}]}
    assert rig["pages"] == [] and rig["failed"] == []
    [call] = stub.calls
    assert call["JobTag"] == "DOC-1"
    assert call["DocumentLocation"] == {"S3Object": {"Bucket": "phi-bucket", "Name": "cases/C1/u31-scan.pdf"}}


def test_corrupt_pdf_falls_through_to_textract_without_raising(rig, monkeypatch):
    """Garbage bytes that aren't a PDF at all: pypdf raises on open → _native_pdf_read returns the
    fall-through sentinel (never crashes start_handler) → Textract owns it."""
    stub = _TextractStub()
    monkeypatch.setattr(handler, "textract", stub)
    monkeypatch.setenv("COMPLETION_SNS_TOPIC_ARN", "arn:aws:sns:us-east-1:0:topic")
    monkeypatch.setenv("TEXTRACT_SNS_ROLE_ARN", "arn:aws:iam::0:role/textract")
    rig["with_bytes"](b"this is not a pdf at all, not even a header")

    result = handler.start_handler(_event("cases/C1/u32-broken.pdf"), None)

    assert result == {"started": [{"documentId": "DOC-1", "jobId": "JOB-1"}]}
    assert rig["pages"] == [] and rig["failed"] == []
    assert len(stub.calls) == 1


def test_native_pdf_hybrid_doc_posts_empty_for_image_pages(rig):
    """A digital doc with a couple of scanned image pages mixed in: the image pages extract to '' and
    post as EMPTY pages (acceptable, noted) — the whole doc is NOT failed for a few image pages. Built
    by interleaving blank (image-only) pages into a real text PDF."""
    from pypdf import PdfWriter, PdfReader as _RR

    text_pdf = _make_text_pdf([
        f"Clinical progress note page {i}. Veteran seen in primary care clinic for chronic conditions "
        f"with real text content documented including vital signs, active problem list, and current "
        f"medications reconciled at this encounter on the dated visit for page {i} of the record."
        for i in range(1, 7)
    ])
    w = PdfWriter()
    for i, pg in enumerate(_RR(io.BytesIO(text_pdf)).pages):
        w.add_page(pg)
        if i in (1, 3):  # interleave a blank (image-only) page after some text pages
            w.add_blank_page(width=200, height=200)
    buf = io.BytesIO()
    w.write(buf)
    rig["with_bytes"](buf.getvalue())

    result = handler.start_handler(_event("cases/C1/u33-hybrid.pdf"), None)

    assert result["method"] == "native_pdf_text"
    assert result["emptyPages"] >= 2          # the 2 interleaved blank pages posted empty
    [(_, pages, count)] = rig["pages"]
    assert count == result["pdfPages"]
    assert len([p for p in pages if p["text"] == ""]) >= 2
    assert any("real text content" in p["text"] for p in pages)  # text pages kept their content


def test_native_pdf_batches_posts_over_2000_pages(rig, monkeypatch):
    """A dump larger than the route's 2,000-entries-per-POST cap is split into batches of
    _PAGES_PER_POST, the TRUE page count sent on EVERY batch. Mock PdfReader so the test stays fast
    (no 2,300-page real PDF)."""

    class _FakePage:
        def __init__(self, n: int):
            self._n = n

        def extract_text(self) -> str:
            return (
                f"Page {self._n} clinical record. Substantive medical content for this page including "
                f"the encounter date, the assessment and plan, the active medication list, and the "
                f"provider signature block, well above the per-page text-layer probe floor for page {self._n}."
            )

    class _FakePages:
        def __init__(self, count: int):
            self._count = count

        def __len__(self) -> int:
            return self._count

        def __getitem__(self, i: int) -> "_FakePage":
            return _FakePage(i)

    class _FakeReader:
        is_encrypted = False

        def __init__(self, _stream):
            self.pages = _FakePages(2300)

    import pypdf
    monkeypatch.setattr(pypdf, "PdfReader", _FakeReader)
    rig["with_bytes"](b"%PDF-1.4 fake header; PdfReader is mocked")

    posts: list = []
    monkeypatch.setattr(handler, "_post_pages_to_api",
                        lambda doc_id, pages, count: posts.append((doc_id, len(pages), count)))

    result = handler.start_handler(_event("cases/C1/u34-huge.pdf"), None)

    assert result["method"] == "native_pdf_text"
    assert result["pdfPages"] == 2300
    # 2300 pages / 1000 per POST = 3 batches (1000, 1000, 300); EVERY batch carries the true 2300 count.
    assert [n for (_, n, _) in posts] == [1000, 1000, 300]
    assert all(count == 2300 for (_, _, count) in posts)
    assert all(doc_id == "DOC-1" for (doc_id, _, _) in posts)


def test_native_pdf_over_size_cap_defers_to_textract(rig, monkeypatch):
    """A PDF larger than MAX_OCR_BYTES is not buffered for a native read — Textract (which streams
    from S3) owns it. The native branch must not even open it."""
    stub = _TextractStub()
    monkeypatch.setattr(handler, "textract", stub)
    monkeypatch.setenv("COMPLETION_SNS_TOPIC_ARN", "arn:aws:sns:us-east-1:0:topic")
    monkeypatch.setenv("TEXTRACT_SNS_ROLE_ARN", "arn:aws:iam::0:role/textract")
    monkeypatch.setattr(handler, "MAX_OCR_BYTES", 16)
    rig["with_bytes"](b"x" * 64)  # over the 16-byte cap

    result = handler.start_handler(_event("cases/C1/u35-oversize.pdf"), None)

    assert result == {"started": [{"documentId": "DOC-1", "jobId": "JOB-1"}]}
    assert rig["pages"] == [] and rig["failed"] == []
    assert len(stub.calls) == 1


# ===== guards around the native branch =====

def test_pdf_still_goes_to_textract_when_text_layer_thin(rig, monkeypatch):
    """The original guard, updated for Layer 1: a PDF with no usable text layer still reaches Textract
    with JobTag=documentId. (A textless single-page PDF stands in for the generic scan.)"""
    stub = _TextractStub()
    monkeypatch.setattr(handler, "textract", stub)
    monkeypatch.setenv("COMPLETION_SNS_TOPIC_ARN", "arn:aws:sns:us-east-1:0:topic")
    monkeypatch.setenv("TEXTRACT_SNS_ROLE_ARN", "arn:aws:iam::0:role/textract")
    rig["with_bytes"](_make_blank_pdf(1))

    result = handler.start_handler(_event("cases/C1/u14-records.pdf"), None)

    assert result == {"started": [{"documentId": "DOC-1", "jobId": "JOB-1"}]}
    assert rig["pages"] == [] and rig["failed"] == []
    [call] = stub.calls
    assert call["DocumentLocation"] == {"S3Object": {"Bucket": "phi-bucket", "Name": "cases/C1/u14-records.pdf"}}
    assert call["JobTag"] == "DOC-1"


def test_delivery_output_is_skipped_before_resolve(rig):
    # cases/<id>/delivery/<artifact>-v<n>.pdf is a GENERATED output (cover memo), not an uploaded
    # record. It must skip BEFORE _resolve_document — otherwise it 404s, raises (key startswith
    # cases/), exhausts retries, and floods the ocr-start DLQ, holding the runaway alarm RED.
    # (aws-cloud-sme audit 2026-06-17; mirrors the /_rendered/ skip.)
    result = handler.start_handler(_event("cases/C1/delivery/cover-memo-v7.pdf"), None)
    assert result.get("skipped") == "delivery_output"
    assert result["started"] == []
    assert rig["pages"] == [] and rig["failed"] == []


def test_has_pages_guard_runs_before_native_read(rig):
    rig["doc"]["hasPages"] = True
    s3 = rig["with_bytes"](b"already read")
    result = handler.start_handler(_event("cases/C1/u15-rerun.txt"), None)
    assert result["skipped"] == "already_has_pages"
    assert rig["pages"] == [] and rig["failed"] == []
    assert s3.calls == []


def test_intake_prefix_is_untouched_by_the_native_branch(rig, monkeypatch):
    sentinel = {"started": [{"intakeS3Key": "intake/abc.txt", "jobId": "J", "jobTag": "t"}]}
    called = []
    monkeypatch.setattr(handler, "_start_intake_ocr", lambda bucket, key: (called.append((bucket, key)), sentinel)[1])
    s3 = rig["with_bytes"](b"should never be fetched")

    result = handler.start_handler(_event("intake/abc.txt"), None)

    assert result is sentinel
    assert called == [("phi-bucket", "intake/abc.txt")]
    assert s3.calls == [] and rig["pages"] == []


def test_oversize_native_file_flags_for_rn(rig, monkeypatch):
    monkeypatch.setattr(handler, "MAX_OCR_BYTES", 16)
    rig["with_bytes"](b"x" * 64)
    result = handler.start_handler(_event("cases/C1/u16-huge.txt"), None)
    assert result["flaggedForRn"] is True
    [(_, status, _, note)] = rig["failed"]
    assert status == "NATIVE_TOO_LARGE"
    assert "summarize manually" in note


def test_extensionless_key_does_not_native_read(rig, monkeypatch):
    stub = _TextractStub()
    monkeypatch.setattr(handler, "textract", stub)
    monkeypatch.setenv("COMPLETION_SNS_TOPIC_ARN", "arn:aws:sns:us-east-1:0:topic")
    monkeypatch.setenv("TEXTRACT_SNS_ROLE_ARN", "arn:aws:iam::0:role/textract")

    handler.start_handler(_event("cases/C1/u17-no-extension"), None)

    assert rig["pages"] == []
    assert len(stub.calls) == 1  # falls through to Textract, as before


# ===== unit coverage of the page splitter =====

def test_build_native_pages_chunks_oversize_text(monkeypatch):
    monkeypatch.setattr(handler, "_MAX_PAGE_CHARS", 10)
    pages = handler._build_native_pages("abcdefghijKLMNOPQRSTuv")
    assert [p["text"] for p in pages] == ["abcdefghij", "KLMNOPQRST", "uv"]
    assert [p["pageNumber"] for p in pages] == [1, 2, 3]


# ===== Package 4a: orphan-race fix — raise-for-retry on unresolvable cases/ keys =====
# The Document row lands a beat AFTER the S3 object (recordDocument race), so the by-s3-key
# lookup can 404 on a perfectly good upload. Returning success silently dropped the file;
# raising lets the Lambda async retry (retryAttempts: 2) re-resolve once the row exists.
# HARD SCOPE GUARD (the highest-risk regression per the plan): ONLY cases/ raises —
# intake/ has no Document by design and must keep returning cleanly.

def test_unresolvable_cases_key_raises_for_async_retry(rig, monkeypatch):
    rig["doc"] = None  # _resolve_document → None: the 404/no-row outcome
    monkeypatch.setattr(handler, "_resolve_document", lambda key: rig["doc"])

    with pytest.raises(RuntimeError, match=r"no resolvable document .*cases/C1/u18-race\.pdf"):
        handler.start_handler(_event("cases/C1/u18-race.pdf"), None)

    # Nothing was posted or flagged — the event must surface to Lambda for the async retry,
    # not get half-processed first.
    assert rig["pages"] == [] and rig["failed"] == []


def test_unresolvable_intake_key_does_not_raise(rig, monkeypatch):
    """Regression guard for the scope mistake: intake/ objects legitimately have no Document
    (parse-at-intake caches to IntakePage) — the orphan-race raise must never reach them. Even
    when the intake ocr-start record fails (the intake no-doc analog), the handler returns
    cleanly so the file falls through to assign-time OCR."""
    rig["doc"] = None
    resolve_calls = []
    monkeypatch.setattr(handler, "_resolve_document", lambda key: (resolve_calls.append(key), rig["doc"])[1])

    def _record_fails(intake_s3_key, job_tag):
        raise RuntimeError("intake ocr-start record failed: 500")

    monkeypatch.setattr(handler, "_post_intake_ocr_start", _record_fails)

    result = handler.start_handler(_event("intake/I1/upload.pdf"), None)  # must NOT raise

    assert result == {"started": []}
    assert resolve_calls == []  # intake/ branches off BEFORE the by-s3-key lookup
    assert rig["pages"] == [] and rig["failed"] == []


def test_unresolvable_non_cases_non_intake_key_keeps_skip_behavior(rig, monkeypatch):
    """A key under any OTHER prefix (defensive: the rule only matches cases/ + intake/, but a
    manual test-invoke or future rule edit shouldn't retry-loop) keeps the old skip-with-success."""
    rig["doc"] = None
    monkeypatch.setattr(handler, "_resolve_document", lambda key: rig["doc"])

    result = handler.start_handler(_event("scratch/manual-test.pdf"), None)  # must NOT raise

    assert result == {"started": []}
    assert rig["pages"] == [] and rig["failed"] == []


def test_event_with_no_bucket_or_key_still_returns_success(rig):
    """The malformed-event early-return sits ABOVE the raise and must keep returning success —
    retrying a bucket-less event can never resolve anything."""
    assert handler.start_handler({"detail": {}}, None) == {"started": []}


def test_resolved_cases_key_is_unaffected_by_the_orphan_raise(rig, monkeypatch):
    """A resolvable Document rides the exact pre-4a path: native .txt reads post pages, no raise."""
    rig["with_bytes"](b"Resolved fine: lumbar strain, chronic.")
    result = handler.start_handler(_event("cases/C1/u19-resolved.txt"), None)
    assert result["started"] == []
    assert result["native"] == "txt"
    assert len(rig["pages"]) == 1 and rig["failed"] == []


# ===== CLAUDE_VISION_DESCRIBE: textless-image auto-describe (dark, default OFF) =====
# A textless clinical photo (e.g. an injured leg) yields 0 chars from Textract + Claude OCR, so it hits
# _handle_unreadable. With the flag ON we make a SECOND Claude vision DESCRIBE call and post the stamped
# description as page text so the readiness char-floor passes; "NO CLINICAL CONTENT" must NOT post; with
# the flag OFF the describe call must never fire (behavior unchanged → file flags for the RN).


class _DescribeUrlopen:
    """Stub of urllib.request.urlopen for the describe call. Captures the request body and returns a
    canned Anthropic /v1/messages response carrying `text` as a single text content block."""

    def __init__(self, text: str):
        self.text = text
        self.requests: list = []

    def __call__(self, req, timeout=None):  # noqa: D401 — mimics urlopen(req, timeout=...)
        self.requests.append(req)
        payload = {"content": [{"type": "text", "text": self.text}], "stop_reason": "end_turn"}
        return _FakeHTTPResponse(json.dumps(payload).encode("utf-8"))


class _FakeHTTPResponse:
    def __init__(self, body: bytes):
        self._body = body
        self.status = 200

    def read(self):
        return self._body

    def __enter__(self):
        return self

    def __exit__(self, *exc):
        return False


@pytest.fixture
def describe_rig(monkeypatch):
    """Wire _handle_unreadable's describe path: a resolvable image document, a stubbed S3 + secret +
    Claude vision call, and capture of every /pages and failed-read-attempt post."""
    state = {"pages": [], "failed": []}
    monkeypatch.setattr(handler, "_post_pages_to_api",
                        lambda doc_id, pages, count: state["pages"].append((doc_id, pages, count)))
    monkeypatch.setattr(handler, "_post_failed_read_attempt",
                        lambda doc_id, status, job_id, error_message=None: state["failed"].append((doc_id, status, job_id, error_message)))
    # describe-call dependencies
    monkeypatch.setattr(handler, "_anthropic_key", lambda: "sk-ant-test")
    monkeypatch.setattr(handler, "_document_source", lambda doc_id: {"s3Key": "cases/C1/u20-leg.png", "contentType": "image/png"})
    monkeypatch.setattr(handler, "_phi_bucket", lambda: "phi-bucket")
    monkeypatch.setattr(handler, "s3", _S3Stub(b"\x89PNG\r\n\x1a\n fake image bytes"))
    # the verbatim Claude OCR path is exercised separately; here it reads nothing so the describe path runs
    monkeypatch.setattr(handler, "_claude_ocr", lambda doc_id: "")
    state["set_urlopen"] = lambda text: monkeypatch.setattr(handler.urllib.request, "urlopen", _DescribeUrlopen(text))
    return state


def test_vision_describe_off_does_not_post_or_call(describe_rig, monkeypatch):
    """Flag OFF (default): _handle_unreadable falls straight through to the RN flag — no describe call,
    no posted page. Behavior is identical to before this change."""
    monkeypatch.delenv("CLAUDE_VISION_DESCRIBE", raising=False)
    called = []
    monkeypatch.setattr(handler.urllib.request, "urlopen",
                        lambda *a, **k: called.append(a) or (_ for _ in ()).throw(AssertionError("describe must not call urlopen when flag off")))

    read = handler._handle_unreadable("DOC-IMG", "EMPTY", "JOB-IMG")

    assert read is False
    assert describe_rig["pages"] == []          # nothing posted as text
    assert len(describe_rig["failed"]) == 1     # flagged for RN exactly as before
    assert called == []


def test_vision_describe_on_posts_stamped_description(describe_rig, monkeypatch):
    """Flag ON + a real description: the description is posted as page text PREFIXED with the AI-visual-
    evidence provenance marker (never reads as OCR'd record text), and the file is NOT flagged for RN."""
    monkeypatch.setenv("CLAUDE_VISION_DESCRIBE", "on")
    describe_rig["set_urlopen"]("Right lower leg with a 6 cm linear surgical scar; mild swelling over the ankle.")

    read = handler._handle_unreadable("DOC-IMG", "EMPTY", "JOB-IMG")

    assert read is True
    [(doc_id, pages, count)] = describe_rig["pages"]
    assert doc_id == "DOC-IMG" and count == 1
    text = pages[0]["text"]
    assert text.startswith(handler._IMAGE_EVIDENCE_PREFIX)  # provenance stamp present
    assert "not OCR text" in text
    assert "surgical scar" in text                          # the description body survived
    assert describe_rig["failed"] == []                     # not dead-ended; letter unblocked


def test_vision_describe_no_clinical_content_does_not_post(describe_rig, monkeypatch):
    """Flag ON but the model returns the exact NO CLINICAL CONTENT sentinel: do NOT post a description —
    leave it as a failed read so the manual path owns it (flagged for RN)."""
    monkeypatch.setenv("CLAUDE_VISION_DESCRIBE", "on")
    describe_rig["set_urlopen"](handler._NO_CLINICAL_CONTENT)

    read = handler._handle_unreadable("DOC-IMG", "EMPTY", "JOB-IMG")

    assert read is False
    assert describe_rig["pages"] == []          # no meaningless description posted
    assert len(describe_rig["failed"]) == 1     # falls through to the RN flag


def test_vision_describe_skips_non_image_media(describe_rig, monkeypatch):
    """Flag ON but the media is a PDF (not png/jpeg): the describe path is image-only, so it must not post
    a description — a textless PDF still flags for the RN."""
    monkeypatch.setenv("CLAUDE_VISION_DESCRIBE", "on")
    monkeypatch.setattr(handler, "_document_source",
                        lambda doc_id: {"s3Key": "cases/C1/u21-scan.pdf", "contentType": "application/pdf"})
    called = []
    monkeypatch.setattr(handler.urllib.request, "urlopen",
                        lambda *a, **k: called.append(a) or (_ for _ in ()).throw(AssertionError("describe must not call urlopen for non-image media")))

    read = handler._handle_unreadable("DOC-PDF", "EMPTY", "JOB-PDF")

    assert read is False
    assert describe_rig["pages"] == []
    assert len(describe_rig["failed"]) == 1
    assert called == []


def test_vision_describe_error_falls_through_to_rn_flag(describe_rig, monkeypatch):
    """Flag ON but the describe Claude call RAISES: _try_image_describe swallows + logs the error and
    returns False, so _handle_unreadable falls through to the RN flag (never blocks the flag path on a
    describe-call failure). The auto-recovery ladder degrades safely to the human last-resort."""
    monkeypatch.setenv("CLAUDE_VISION_DESCRIBE", "on")
    monkeypatch.setattr(handler.urllib.request, "urlopen",
                        lambda *a, **k: (_ for _ in ()).throw(RuntimeError("anthropic 529 overloaded")))

    read = handler._handle_unreadable("DOC-IMG", "EMPTY", "JOB-IMG")

    assert read is False
    assert describe_rig["pages"] == []          # nothing posted on a failed describe call
    assert len(describe_rig["failed"]) == 1     # safely flagged for the RN (last resort)


def test_vision_describe_on_with_text_ocr_never_reaches_describe(describe_rig, monkeypatch):
    """Flag ON, but the verbatim Claude OCR produced real text: that text is posted as-is and the describe
    path never runs (the describe call only fires when OCR yielded nothing)."""
    monkeypatch.setenv("CLAUDE_VISION_DESCRIBE", "on")
    monkeypatch.setattr(handler, "_claude_ocr", lambda doc_id: "AHI 36.4 events/hour. Severe OSA.")
    called = []
    monkeypatch.setattr(handler.urllib.request, "urlopen",
                        lambda *a, **k: called.append(a) or (_ for _ in ()).throw(AssertionError("describe must not run when OCR read real text")))

    read = handler._handle_unreadable("DOC-IMG", "EMPTY", "JOB-IMG")

    assert read is True
    [(_, pages, _)] = describe_rig["pages"]
    assert pages[0]["text"] == "AHI 36.4 events/hour. Severe OSA."   # verbatim OCR, NOT stamped
    assert not pages[0]["text"].startswith(handler._IMAGE_EVIDENCE_PREFIX)
    assert called == []
